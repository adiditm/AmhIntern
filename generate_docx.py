from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Proposal Penawaran Fitur Chat AI', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Intro
doc.add_paragraph('Berikut adalah daftar paket penawaran untuk integrasi Chat AI pada platform AMHIntern. Paket disusun secara bertingkat berdasarkan kelengkapan fitur dan kemandirian manajemen sistem.')

# Paket Basic
doc.add_heading('1. Basic - Rp 2.000.000', level=1)
doc.add_paragraph('Fokus pada asisten publik untuk melayani calon pendaftar/pengunjung.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Halaman Login (Publik): ').bold = True
p.add_run('Asisten virtual untuk menjawab pertanyaan umum seputar program, pendaftaran, dan informasi kontak.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Basis Pengetahuan (Statis): ').bold = True
p.add_run('Informasi diberikan melalui instruksi sistem (system prompt) yang ditanam di dalam kode.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('UI Widget: ').bold = True
p.add_run('Desain chat widget standar di pojok layar halaman login.')

# Paket Standar
doc.add_heading('2. Standar - Rp 3.000.000 (Sudah termasuk Basic)', level=1)
doc.add_paragraph('Fokus pada pelayanan internal dengan sistem basis pengetahuan (RAG) yang dikelola oleh tim Developer.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Halaman Seller & Jamaah: ').bold = True
p.add_run('Chat AI khusus di dashboard seller dan jamaah. AI secara otomatis mengenali peran user dan menyesuaikan jawabannya.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Basis Pengetahuan RAG (Managed Service): ').bold = True
p.add_run('Bot sudah mampu menjawab berdasarkan dokumen SOP atau panduan khusus perusahaan. Namun, proses upload dan pembaruan dokumen dilakukan secara manual oleh tim Developer dari sisi backend.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Biaya Maintenance: ').bold = True
p.add_run('Karena pembaruan dokumen dibantu oleh developer, dapat dikenakan biaya maintenance/update ringan setiap kali ada perubahan dokumen, atau opsi biaya bulanan tetap sesuai kesepakatan (SLA / Service Level Agreement menyesuaikan).')

# Paket Premium
doc.add_heading('3. Premium - Rp 5.000.000 (Sudah termasuk Basic & Standar)', level=1)
doc.add_paragraph('Solusi AI menyeluruh dengan kemandirian penuh untuk Admin dalam mengelola sistem basis pengetahuan (Self-Service RAG).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Halaman Pebisnis (Network Assistant): ').bold = True
p.add_run('Tambahan asisten AI khusus member/upline untuk panduan jaringan downline dan skema komisi.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Halaman Admin (Self-Service RAG Upload): ').bold = True
p.add_run('Admin klien memiliki menu UI khusus di dashboard untuk mengunggah file (PDF/Word) secara mandiri kapan saja. Sistem otomatis memproses dokumen tanpa perlu bantuan developer. Klien bebas mengalokasikan dokumen spesifik untuk tiap role (Seller/Pebisnis/Jamaah).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Analitik Percakapan: ').bold = True
p.add_run('Admin dapat melihat ringkasan topik/pertanyaan yang paling sering diajukan pengguna.')


doc.add_paragraph('\n*Catatan Tambahan: Biaya operasional layanan API dari pihak ketiga (OpenAI / Google Gemini) menggunakan sistem pay-as-you-go menyesuaikan volume penggunaan bulanan. Biaya di atas merupakan biaya implementasi/pengembangan awal (One-time cost).', style='Body Text')

doc.save('/Users/didit/prx/AMHIntern/chat_ai.docx')
